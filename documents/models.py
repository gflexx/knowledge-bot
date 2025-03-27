from django.core.exceptions import ValidationError
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from django.db.models.signals import pre_delete
from django.dispatch import receiver
from django.apps import apps
from django.db import models
import shutil
import os

from .llm import create_knowledge_base, add_document_to_vector_store, get_document_vector_store_dir

def validate_document_type(value):
    """Ensure only PDF or DOCX files are uploaded."""
    ext = os.path.splitext(value.name)[1].lower()
    if ext not in ['.pdf', '.docx']:
        raise ValidationError("Only PDF and DOCX files are allowed.")


class Document(models.Model):
    title = models.CharField(max_length=252)
    file = models.FileField(
        upload_to="documents/",
        validators=[validate_document_type]
    )
    priority = models.BooleanField(default=False)
    creation_time = models.DateTimeField(auto_now_add=True)
    last_updated_time = models.DateTimeField(auto_now=True)

    class Meta:
        ordering = ['-creation_time',]

    def __str__(self):
        return self.title
    

    def save(self, *args, **kwargs):
        """
        on save, add to vector store or initialize it
        """
        super().save(*args, **kwargs) 

        # get vector store create if empty else add doc
        documents_config = apps.get_app_config('documents')
        if documents_config.vector_stores is None:
            documents_config.vector_stores = create_knowledge_base()
            print("New vector store created from documents...")

        else:
            print(f"Adding {self.title} to vector store...")
            add_document_to_vector_store(self)


    def extract_text(self):
        """
        extract text from the document based on file type
        """
        ext = os.path.splitext(self.file.name)[1].lower()
        extracted_pages = []

        if ext == ".pdf":
            reader = PdfReader(self.file)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    page_text += f" {self.title}"
                    extracted_pages.append(
                        {"text": page_text, "page": page_num}
                    )

        elif ext == ".docx":
            doc = DocxDocument(self.file)
            current_page = 1
            char_count = 0
            chars_per_page = 3000  # approximate character count per page
            page_text = ""

            for para in doc.paragraphs:
                if para.text.strip():
                    page_text += para.text + " "
                    page_text += f" {self.title}"
                    char_count += len(para.text)

                # check for manual page breaks
                if para.text.strip() == "":
                    if "page-break-before" in para._element.xml or "w:br w:type='page'" in para._element.xml:
                        extracted_pages.append(
                            {
                                "text": page_text.strip(), 
                                "page": current_page
                            }
                        )
                        current_page += 1
                        page_text = ""
                        char_count = 0

                # approximate page breaks based on character count
                elif char_count >= chars_per_page:
                    page_text += f" {self.title}"
                    extracted_pages.append(
                        {
                            "text": page_text.strip(), 
                            "page": current_page
                        }
                    )
                    current_page += 1
                    page_text = ""
                    char_count = 0

                # Add last collected page text
                if page_text.strip():
                    page_text += f" {self.title}"
                    extracted_pages.append(
                        {
                            "text": page_text.strip(), 
                            "page": current_page
                        }
                    )

        return extracted_pages
    
@receiver(pre_delete, sender=Document)
def delete_vector_store(sender, instance, **kwargs):
    documents_config = apps.get_app_config("documents")
    vector_stores = documents_config.vector_stores 
    doc_vector_store_dir = get_document_vector_store_dir(instance.id)

    if os.path.exists(doc_vector_store_dir):
        shutil.rmtree(doc_vector_store_dir) 

    if instance.id in vector_stores:
        del vector_stores[instance.id]

    print(f"Deleted vector store for document {instance.id}: {instance.title}")